"""Utilities for converting Markdown content to DOCX."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import List

from docx import Document


def _add_runs_with_emphasis(paragraph, text: str) -> None:
	"""Add text runs with markdown emphasis (bold, code) to a paragraph."""
	if not text:
		return
	# Minimal inline markdown support: **bold**, __bold__, `code`
	parts = re.split(r"(\*\*.+?\*\*|__.+?__|`.+?`)", text)
	for part in parts:
		if not part:
			continue
		if (part.startswith("**") and part.endswith("**") and len(part) >= 4):
			run = paragraph.add_run(part[2:-2])
			run.bold = True
		elif (part.startswith("__") and part.endswith("__") and len(part) >= 4):
			run = paragraph.add_run(part[2:-2])
			run.bold = True
		elif (part.startswith("`") and part.endswith("`") and len(part) >= 2):
			# Render inline code as a plain run for now
			paragraph.add_run(part[1:-1])
		else:
			paragraph.add_run(part)


def _add_table_to_docx(document, table_rows: List[List[str]]) -> None:
	"""Convert markdown table rows to a DOCX table."""
	if not table_rows:
		return
	
	# Determine number of columns from the first row
	num_cols = len(table_rows[0]) if table_rows else 0
	if num_cols == 0:
		return
	
	try:
		# Create table with appropriate dimensions
		table = document.add_table(rows=len(table_rows), cols=num_cols)
		# Try to set a nice table style, but fall back if it doesn't exist
		try:
			table.style = "Light Grid Accent 1"
		except Exception:
			# Style doesn't exist, use default
			pass
		
		# Fill in the table cells
		for row_idx, row_data in enumerate(table_rows):
			row = table.rows[row_idx]
			for col_idx in range(num_cols):
				cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
				# Remove markdown escape characters like \#
				cell_text = cell_text.replace("\\#", "#")
				cell = row.cells[col_idx]
				cell_para = cell.paragraphs[0]
				# Clear default paragraph text
				cell_para.clear()
				_add_runs_with_emphasis(cell_para, cell_text)
				
				# Make header row bold (first row)
				if row_idx == 0:
					for run in cell_para.runs:
						run.bold = True
	except Exception as e:
		# If table creation fails, add as plain text
		para = document.add_paragraph(f"[Table conversion failed: {e}]")
		for row in table_rows:
			para = document.add_paragraph(" | ".join(row))


def markdown_to_docx_bytes(content: str) -> io.BytesIO:
	"""Convert a subset of Markdown into DOCX bytes."""
	document = Document()
	in_code_block = False
	code_buffer: List[str] = []

	lines = content.splitlines()
	i = 0
	table_buffer: List[List[str]] = []
	in_table = False
	
	while i < len(lines):
		raw_line = lines[i]
		line = raw_line.rstrip("\n")
		stripped = line.strip()

		# Check if this is a table row (starts with |)
		# Also handle tables that might not end with | (some markdown variants)
		is_table_row = stripped.startswith("|")
		if is_table_row:
			# Parse table row - handle both |...| and |... (without trailing |)
			if stripped.endswith("|"):
				cells = [cell.strip() for cell in stripped[1:-1].split("|")]
			else:
				cells = [cell.strip() for cell in stripped[1:].split("|")]
			
			# Filter out empty cells from splitting
			cells = [c for c in cells if c or len(cells) > 1]  # Keep at least one cell even if empty
			
			# Skip separator rows (like |---|---| or |:----|:----|)
			is_separator = all(
				cell.replace("-", "").replace(":", "").replace(" ", "").strip() == "" 
				for cell in cells
			)
			if is_separator:
				i += 1
				continue
			
			# Only treat as table if we have at least 2 columns
			if len(cells) >= 2:
				table_buffer.append(cells)
				in_table = True
				i += 1
				continue
		
		# If we were in a table and hit a non-table line, process the table
		if in_table and table_buffer:
			_add_table_to_docx(document, table_buffer)
			table_buffer.clear()
			in_table = False
			# Don't increment i, process this line normally
			continue

		if stripped.startswith("```"):
			if in_code_block:
				paragraph = document.add_paragraph("\n".join(code_buffer) if code_buffer else "")
				paragraph.style = "Intense Quote"
				code_buffer.clear()
				in_code_block = False
			else:
				in_code_block = True
			i += 1
			continue

		if in_code_block:
			code_buffer.append(line)
			i += 1
			continue

		if not stripped:
			document.add_paragraph("")
			i += 1
			continue

		if stripped.startswith("#"):
			level = len(stripped) - len(stripped.lstrip("#"))
			text = stripped[level:].strip()
			level = max(1, min(level, 4))
			heading_para = document.add_heading("", level=level)
			_add_runs_with_emphasis(heading_para, text or "")
			i += 1
			continue

		if stripped.startswith(('- ', '* ')):
			bullet_text = stripped[2:].strip()
			bullet_para = document.add_paragraph(style="List Bullet")
			_add_runs_with_emphasis(bullet_para, bullet_text)
			i += 1
			continue

		if stripped.startswith(">"):
			quote_para = document.add_paragraph("")
			quote_para.style = "Intense Quote"
			_add_runs_with_emphasis(quote_para, stripped[1:].strip())
			i += 1
			continue

		body_para = document.add_paragraph("")
		_add_runs_with_emphasis(body_para, stripped)
		i += 1
	
	# Handle table at end of document
	if in_table and table_buffer:
		_add_table_to_docx(document, table_buffer)

	if code_buffer:
		paragraph = document.add_paragraph("\n".join(code_buffer))
		paragraph.style = "Intense Quote"

	try:
		buffer = io.BytesIO()
		document.save(buffer)
		buffer.seek(0)
		return buffer
	except Exception as e:
		# If saving fails, return empty buffer with error message
		error_doc = Document()
		error_doc.add_paragraph(f"Error converting markdown to DOCX: {e}")
		buffer = io.BytesIO()
		error_doc.save(buffer)
		buffer.seek(0)
		return buffer


def save_markdown_as_docx(content: str, output_path: Path) -> None:
	"""Render Markdown to a DOCX file at output_path."""
	try:
		buf = markdown_to_docx_bytes(content)
		output_path.parent.mkdir(parents=True, exist_ok=True)
		with open(output_path, "wb") as f:
			data = buf.read()
			if len(data) == 0:
				# If buffer is empty, create a minimal document
				error_doc = Document()
				error_doc.add_paragraph("Error: Generated DOCX was empty. Original markdown may have issues.")
				error_doc.save(output_path)
			else:
				f.write(data)
	except Exception as e:
		# Create error document if conversion fails
		error_doc = Document()
		error_doc.add_paragraph(f"Error converting markdown to DOCX: {e}")
		error_doc.save(output_path)


